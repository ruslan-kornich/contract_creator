import docx
from django.conf import settings


def contract_create(data: dict) -> str:
    doc = docx.Document(settings.BLANK_PATH)

    for i in data:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                style = doc.styles["Normal"]
                font = style.font
                font.name = "Times New Roman"
                font.size = docx.shared.Pt(12)
                p.text = p.text.replace(i, data[i])
    for j in data:
        for table in doc.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.find(j) >= 0:
                            paragraph.text = paragraph.text.replace(j, data[j])
    contract_path = settings.CONTRACT_PATH + f'{data.get("short_name")}.docx'
    doc.save(contract_path)
    return contract_path
